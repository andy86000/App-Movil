from docx import Document

def exportar_a_word(texto, nombre_archivo="outputs/resumen_legal.docx"):
    doc = Document()
    doc.add_heading("Resumen del Documento Legal", 0)
    doc.add_paragraph(texto)
    doc.save(nombre_archivo)
